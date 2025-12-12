import streamlit as st
import google.generativeai as genai
import tempfile
import os
import time
from pathlib import Path
from docx import Document
from io import BytesIO

# --- 頁面設定 ---
st.set_page_config(
    page_title="Landing Page 策略 Gemini 3.0 工作台",
    page_icon="🧩",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- CSS 優化 ---
st.markdown("""
    <style>
    .stTextArea textarea { font-size: 14px; }
    .stButton button { width: 100%; border-radius: 8px; font-weight: bold; }
    .block-container { padding-top: 2rem; padding-bottom: 2rem; }
    div[data-testid="stExpander"] div[role="button"] p { font-size: 1.1rem; font-weight: 600; }
    </style>
""", unsafe_allow_html=True)

# =========================
# Word 處理函式
# =========================

def extract_text_from_docx(file_path: str) -> str:
    """從 docx 檔案路徑提取文字"""
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        return '\n\n'.join(full_text)
    except Exception as e:
        st.error(f"Word 讀取失敗: {e}")
        return ""

def create_docx_from_markdown(markdown_text: str) -> BytesIO:
    """將 Markdown 文字轉換為 Word 檔案物件 (BytesIO)"""
    doc = Document()
    
    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            clean_line = line.replace('**', '').replace('__', '')
            doc.add_paragraph(clean_line)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =========================
# Gemini 核心功能函式
# =========================

def configure_gemini(api_key: str) -> bool:
    """設定 API Key"""
    if not api_key:
        st.error("❌ 請先在側邊欄輸入 Google Gemini API Key")
        return False
    try:
        genai.configure(api_key=api_key)
        return True
    except Exception as e:
        st.error(f"API Key 設定失敗: {e}")
        return False

def process_uploaded_file(uploaded_file):
    """
    處理上傳檔案：
    - docx 會轉 txt 再上傳
    - 其他格式直接上傳
    - 支援圖片、PDF、影片、文字檔
    回傳: Gemini File Object 或 None
    """
    if uploaded_file is None:
        return None
    
    try:
        suffix = Path(uploaded_file.name).suffix.lower()
        tmp_path = ""
        display_name = uploaded_file.name

        # 建立暫存檔案
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name

        # 特殊處理 docx：先抽文字轉 txt，因為 Gemini API 對純文字檔支援度極佳
        if suffix == '.docx':
            text_content = extract_text_from_docx(tmp_path)
            os.remove(tmp_path)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8') as txt_tmp:
                txt_tmp.write(text_content)
                tmp_path = txt_tmp.name
                display_name = f"{uploaded_file.name}.txt"

        # 上傳到 Gemini File API
        with st.spinner(f"正在上傳並處理檔案: {uploaded_file.name} ..."):
            gemini_file = genai.upload_file(path=tmp_path, display_name=display_name)
            
            # 等待處理完成（加上 timeout）
            max_wait = 60  # 最多等 60 秒
            elapsed = 0
            while getattr(gemini_file, "state", None) and gemini_file.state.name == "PROCESSING" and elapsed < max_wait:
                time.sleep(1)
                elapsed += 1
                gemini_file = genai.get_file(gemini_file.name)
            
            if not getattr(gemini_file, "state", None) or gemini_file.state.name == "FAILED" or elapsed >= max_wait:
                st.error(f"檔案 {uploaded_file.name} 處理失敗或逾時。")
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
                return None
                
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
            
        return gemini_file

    except Exception as e:
        st.error(f"上傳錯誤 ({uploaded_file.name}): {e}")
        return None

def generate_content_stream(model_name: str, prompt: str, files=None) -> str | None:
    """呼叫 Gemini API 生成內容"""
    try:
        model = genai.GenerativeModel(model_name)
        if files is None:
            files = []
        
        content_parts = [prompt]
        if files:
            content_parts.extend(files)
            
        with st.spinner(f"正在使用 {model_name} 模型進行深度運算中..."):
            response = model.generate_content(
                content_parts,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.7,
                )
            )
        return response.text
    except Exception as e:
        st.error(f"生成錯誤: {e}")
        return None

# =========================
# Session State 初始化
# =========================

if 'step1_result' not in st.session_state:
    st.session_state.step1_result = ""
if 'step2_result' not in st.session_state:
    st.session_state.step2_result = ""
if 'step3_result' not in st.session_state:
    st.session_state.step3_result = ""

# =========================
# 側邊欄設定
# =========================

with st.sidebar:
    st.header("⚙️ 系統設定")
    api_key = st.text_input("輸入 Gemini API Key", type="password", help="請輸入您的 Google AI Studio API Key")
    
    st.markdown("### 🧠 模型選擇")
    # 更新為使用者指定的模型清單
    model_options = [
        "gemini-3-pro",
        "gemini-2.5-pro",
        "gemini-3-pro-preview",
        "gemini-2.5-flash"
    ]
    selected_model = st.selectbox("使用模型", model_options, index=0)
    
    st.markdown("---")
    st.info(f"當前優先使用: **{selected_model}**")
    
    if st.button("重置所有分析", type="secondary"):
        st.session_state.step1_result = ""
        st.session_state.step2_result = ""
        st.session_state.step3_result = ""
        st.rerun()

# =========================
# 主標題
# =========================

st.title("🧩 Landing Page 策略 Gemini 3.0 工作台")
st.markdown("### 競品拆解 → 我方情境診斷 → LP 結構與文案產出")
st.markdown("---")

tab1, tab2, tab3 = st.tabs([
    "Step 1: 競品 LP 逆向工程",
    "Step 2: 我方情境診斷（3 模式）",
    "Step 3: LP 產出（Full / Partial）"
])

# ==========================================
# Step 1: 競品 LP 深度分析
# ==========================================
with tab1:
    st.subheader("Step 1: 競品 Landing Page 戰略拆解")
    st.markdown("**目標**：上傳競品 LP（截圖 / PDF / Word / 文字說明），產出完整逆向工程報告。")
    
    col1, col2 = st.columns([1, 1])
    with col1:
        competitor_files = st.file_uploader(
            "上傳競品素材 (可多選: Word/PDF/圖/影/文字)", 
            accept_multiple_files=True,
            type=['docx', 'png', 'jpg', 'jpeg', 'pdf', 'mp4', 'txt', 'csv'],
            key="s1_files"
        )
    with col2:
        competitor_text = st.text_area("貼上競品 LP 連結 / 文案 / 結構描述 (選填)", height=180)

    if st.button("🚀 執行 Step 1 競品分析", type="primary", key="btn_s1"):
        if configure_gemini(api_key):
            gemini_files_s1 = []
            if competitor_files:
                for f in competitor_files:
                    g_file = process_uploaded_file(f)
                    if g_file:
                        gemini_files_s1.append(g_file)
            
            prompt_s1 = f"""
# Role
你是一位資深 CRO 顧問與 Landing Page 策略設計師。

# 任務目標
請針對我提供的【競品 Landing Page 資料】進行「結構化的逆向工程」，輸出一份《競品 Landing Page 戰略拆解報告》。

競品資料來源包含：
1. 我上傳的檔案與截圖（請一併納入分析）
2. 我補充貼上的文字說明或網址：
{competitor_text}

# 分析觀點
請聚焦在「如何說服訪客完成轉換」，從以下角度拆解：

1. **定位與主張 (Positioning & Core Promise)**
2. **頁面結構與資訊架構 (Information Architecture)**
3. **首屏設計 (Above-the-fold Strategy)**
4. **文案語氣與說服手法 (Copy & Persuasion Technique)**
5. **視覺與互動策略 (Visual & Interaction)**
6. **轉換路徑設計 (Conversion Path)**
7. **關鍵說服模組庫 (Reusable Blocks)**
8. **整體策略總結 (Strategic Summary)**

請使用 Markdown 輸出，最後附上一個小節：
**《可直接套用的設計原則清單》**，條列 5–10 點即可。
"""
            result = generate_content_stream(selected_model, prompt_s1, gemini_files_s1)
            if result:
                st.session_state.step1_result = result
                st.success("Step 1 競品分析完成！")

    if st.session_state.step1_result:
        st.markdown("---")
        st.markdown("### 📝 Step 1 分析結果")
        st.markdown(st.session_state.step1_result)
        
        docx_file = create_docx_from_markdown(st.session_state.step1_result)
        st.download_button(
            label="📥 下載 Step 1 報告 (.docx)",
            data=docx_file,
            file_name="Step1_Competitor_LP_Analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ==========================================
# Step 2: 我方情境診斷（3 模式）
# ==========================================
with tab2:
    st.subheader("Step 2: 我方 Landing Page 情境診斷（3 種模式）")
    
    if not st.session_state.step1_result:
        st.warning("⚠️ 請先完成 Step 1 競品分析。")
    else:
        st.markdown("**目標**：依目前情境選擇模式，產出對應的診斷報告。")
        
        # 模式選擇
        step2_mode = st.selectbox(
            "選擇 Step 2 模式",
            [
                "客戶沒有頁面（No Page）",
                "客戶有頁面但定位不清楚（Unclear）",
                "客戶有頁面且可直接比對（Normal）"
            ],
            index=2
        )
        
        st.markdown("#### 客戶基本設定")
        col_a, col_b, col_c = st.columns(3)
        with col_a:
            client_proposition = st.text_input("客戶主訴求（value proposition）", placeholder="例：讓中小企業 3 天內導入 AI 自動化")
        with col_b:
            client_audience = st.text_input("目標受眾（audience）", placeholder="例：台灣中小企業主 / 行銷主管")
        with col_c:
            client_cta = st.text_input("主要 CTA", placeholder="例：預約諮詢 / 加入 Line / 開始試用")
        
        st.markdown("#### 我方現有資料 / 備註")
        st.info("💡 即使選擇「沒有頁面」，也可以在此上傳產品介紹、客戶訪談紀錄或相關文件供 AI 參考。")
        col1, col2 = st.columns([1, 1])
        with col1:
            our_files = st.file_uploader(
                "上傳文件 (產品資料/現有LP截圖/相關素材)", 
                accept_multiple_files=True,
                type=['docx', 'png', 'jpg', 'jpeg', 'pdf', 'mp4', 'txt'],
                key="s2_files"
            )
        with col2:
            our_text = st.text_area(
                "貼上目前頁面文案、架構描述、或你給客戶的備註說明",
                height=200
            )

        if st.button("🚀 執行 Step 2 情境診斷", type="primary", key="btn_s2"):
            if configure_gemini(api_key):
                gemini_files_s2 = []
                
                # --- 修正邏輯：只要有上傳檔案就處理，不限制模式 ---
                if our_files:
                    for f in our_files:
                        g_file = process_uploaded_file(f)
                        if g_file:
                            gemini_files_s2.append(g_file)
                # ---------------------------------------------
                
                # 根據模式選擇 prompt
                if step2_mode == "客戶沒有頁面（No Page）":
                    prompt_s2 = f"""
# Case: 客戶目前沒有 Landing Page。

# 客戶基本資訊
- 主訴求（value proposition）：{client_proposition}
- 目標受眾（audience）：{client_audience}
- 主要 CTA：{client_cta}

# 客戶補充資訊與文件
請參考我上傳的文件（如有，包含產品資料、簡報等）以及以下文字說明：
{our_text}

# 競品洞察背景
以下為 Step 1 的競品拆解摘要：
{st.session_state.step1_result}

# 你的任務
請作為資深 CRO 顧問，協助我為「沒有頁面」的客戶建立一份《Landing Page 初版定位與建議結構》。
請參考我提供的補充文件來萃取產品優勢與痛點。

請輸出：

1. **核心定位摘要**
2. **建議的頁面模組架構（從零開始）**
3. **風險與常見錯誤**
4. **可直接交給 Step 3 的「草稿訊息」：
   - Hero headline＋subheadline
   - 3 個痛點（Problem）
   - 3 個利益點（Benefit）
   - 推薦 CTA 文案
   - 初步信任元素

請用 Markdown 條列清晰輸出。
"""
                elif step2_mode == "客戶有頁面但定位不清楚（Unclear）":
                    prompt_s2 = f"""
# Case: 客戶有頁面，但主訴求與受眾定位不清楚。

# 客戶基本資訊
- 客戶主訴求（value proposition）：{client_proposition}
- 客戶受眾（audience）：{client_audience}
- 主要 CTA：{client_cta}

# 客戶現有頁面資料與補充文件
{our_text}
(請同時參考附件檔案，可能是產品文件或現有頁面截圖)

# 競品背景
Step 1 競品拆解摘要如下：
{st.session_state.step1_result}

# 你的任務
請你站在「CRO 顧問」的角度，協助我進行《定位校正 + 頁面差異診斷》。

請輸出：

1. **頁面與主訴求／受眾是否一致？**
2. **缺失清單（Missing Elements）**
3. **優先修正清單（Immediate Fixes）**
4. **定位校正後的核心訊息草稿：
   - Hero headline + subheadline
   - 受眾描述
   - 3 點利益點
   - 新版 CTA 建議

請以 Markdown 輸出。
"""
                else:  # Normal
                    prompt_s2 = f"""
# Case: 客戶已有頁面，定位清楚，可直接進行差異分析。

# 客戶資訊
- 主訴求：{client_proposition}
- 受眾：{client_audience}
- CTA：{client_cta}

# 客戶現有頁面內容與補充文件
{our_text}
(請同時參考附件檔案)

# 競品拆解摘要
{st.session_state.step1_result}

# 你的任務
請基於以上資訊，產出《我方 Landing Page 差異分析報告》。

請包含：

1. **現況摘要（3–5 句）**
2. **頁面結構對照與差異（Structure vs Competitors）**
3. **訊息缺口（Message Gap）**
4. **轉換阻力（Friction Points）**
5. **高優先項（3–5 個必改項目）**
6. **可放大的優勢（Strengths to Amplify）**

請以 Markdown 輸出。
"""

                result = generate_content_stream(selected_model, prompt_s2, gemini_files_s2)
                if result:
                    st.session_state.step2_result = result
                    st.success("Step 2 情境診斷完成！")

        if st.session_state.step2_result:
            st.markdown("---")
            st.markdown("### 📝 Step 2 診斷結果")
            st.markdown(st.session_state.step2_result)
            
            docx_file_s2 = create_docx_from_markdown(st.session_state.step2_result)
            st.download_button(
                label="📥 下載 Step 2 報告 (.docx)",
                data=docx_file_s2,
                file_name="Step2_LP_Diagnosis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ==========================================
# Step 3: LP 產出（Full / Partial）
# ==========================================
with tab3:
    st.subheader("Step 3: Landing Page 產出（全重建 / 部分優化）")
    
    if not st.session_state.step2_result:
        st.warning("⚠️ 請先完成 Step 1 與 Step 2。")
    else:
        st.markdown("**目標**：依目前狀況選擇是「全面重建」還是「部分優化」。")
        
        step3_mode = st.selectbox(
            "選擇 Step 3 模式",
            [
                "全面重建（Full Rebuild）",
                "部分優化（Partial Optimization）"
            ],
            index=1
        )
        
        st.markdown("#### 產出設定")
        col1, col2 = st.columns(2)
        with col1:
            target_action = st.text_input("主要 CTA（target_action）", value="預約諮詢")
            target_audience = st.text_input("目標受眾說明（target_audience）", value="台灣中小企業主 / 行銷主管")
        with col2:
            extra_constraints = st.text_area(
                "額外限制或品牌規範（extra_constraints，可留白）",
                height=120,
                placeholder="例如：視覺風格需維持現有模板、不改動頁尾、文案語氣偏專業冷靜…"
            )
        
        additional_req = st.text_input(
            "額外要求（可留白）",
            value="請優先產出可以在 1–3 天內實際實作的版本。"
        )
        
        example_file = st.file_uploader(
            "上傳範例文件（選填：想模仿的 LP/版型）", 
            type=['docx', 'pdf', 'jpg', 'png', 'txt'],
            key="s3_example"
        )

        if st.button("🚀 生成 Step 3 LP 產出", type="primary", key="btn_s3"):
            if configure_gemini(api_key):
                gemini_files_s3 = []
                format_instruction = "請使用標準的 Markdown 條列與區塊標題格式輸出。"
                
                if example_file:
                    g_file = process_uploaded_file(example_file)
                    if g_file:
                        gemini_files_s3.append(g_file)
                        format_instruction = "🚨 **格式嚴格要求**：請盡可能模仿附件檔案的「區塊結構」與「欄位架構」，但內容以我方產品為主。"
                
                if step3_mode == "全面重建（Full Rebuild）":
                    prompt_s3 = f"""
# Context
Step 1 - 競品拆解重點：
{st.session_state.step1_result}

Step 2 - 我方情境診斷與建議：
{st.session_state.step2_result}

# 任務目標：全面重建 Landing Page
請你以「資深 CRO 顧問 + Landing Page 產品設計師」的身份，
根據上述 Context，重新打造一個 *完整重構版* Landing Page。

本次目標：
- 目標 CTA：{target_action}
- 核心受眾：{target_audience}
- 額外限制／品牌要求：
{extra_constraints}

# 請輸出《Landing Page 全面重建規格書》，內容需包含：

1. **頁面定位摘要**
2. **全新結構總覽 (New Page Architecture)**
3. **逐區塊規格（Section-by-Section Specification）**
4. **技術備註（For Designer & Front-end）**
5. **A/B Test 起手式建議**

# 格式要求
{format_instruction}

# 額外要求
{additional_req}
"""
                else:
                    prompt_s3 = f"""
# Context
Step 1 - 競品拆解重點：
{st.session_state.step1_result}

Step 2 - 我方情境診斷與 Gap：
{st.session_state.step2_result}

# 任務目標：在有限資源下優化 LP
假設目前「人力與時間有限」，無法全面重做 Landing Page。

本次目標：
- 主要 CTA：{target_action}
- 核心受眾：{target_audience}
- 限制條件：
{extra_constraints}

請你作為資深 CRO 顧問，輸出《Landing Page 有限資源優化指南》，內容包含：

1. 優化優先順序（High / Medium / Low）,每個需優化項目請幫我舉例說明
2. 優先區塊局部改寫（含局部修正版文案）
3. 不可動區塊的建議調整方式
4. 3–5 個「在 1 小時內就能改」的快速提升項目
5. 可延後或下一版再改的 nice-to-have 項目

請使用 Markdown 條列輸出。
{format_instruction}

# 額外要求
{additional_req}
"""

                result = generate_content_stream(selected_model, prompt_s3, gemini_files_s3)
                if result:
                    st.session_state.step3_result = result
                    st.success("Step 3 LP 產出完成！")

        if st.session_state.step3_result:
            st.markdown("---")
            st.markdown("### 🎨 Step 3 產出結果")
            st.markdown(st.session_state.step3_result)
            
            docx_file_s3 = create_docx_from_markdown(st.session_state.step3_result)
            st.download_button(
                label="📥 下載 Step 3 LP 規格檔 (.docx)",
                data=docx_file_s3,
                file_name="Step3_LP_Output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Footer
st.markdown("---")
st.caption(f"Powered by Google {selected_model} | LP Strategic Toolkit")
